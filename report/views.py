from django.http import HttpResponse, JsonResponse
from django.shortcuts import redirect, render
import json

from data_analysis.utils.data_analysis_utils import guardar_resultado_en_sesion

import base64
from io import BytesIO
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 
import requests 

# Aquí están las vistas de todos los reportes que se pueden generar en el sistema
# Las views reciben los datos a través de la sesión y los renderizan en las plantillas correspondientes.
# Cada tipo de analisis tiene su propia vista, y se encarga de recibir los datos del usuario, procesarlos y enviarlos a la vista de reporte.
# También se manejan aquí la exportación de los reportes a un documento Word y el manejo de las imágenes.

# Se incluyen las vistas: reporte_regresion_linear, reporte_regresion_logistica, reporte_tendencia_central, reporte_variabilidad, 
# reporte_distribucion_frecuencias, reporte_correlacion_pearson, reporte_correlacion_spearman, reporte_kmeans_clustering, reporte_arbol_decision_clasificacion, 
# reporte_intervalo_confianza, reporte_prueba_hipotesis, reporte_perceptron, exportar_reporte


@csrf_exempt
@require_http_methods(["POST"])
def exportar_reporte(request):
    try:
        data = json.loads(request.body.decode('utf-8'))

        doc = Document()
        doc.add_heading(data['titulo'], 0)
        doc.add_paragraph(data['descripcion'])

        # Manejo de las tablas
        for tabla_obj in data['tablas']:
            try:
                tabla = tabla_obj['filas']
                titulo_tabla = tabla_obj['titulo']
                if tabla:

                    # título de la tabla como un párrafo antes de la tabla
                    doc.add_paragraph(titulo_tabla, style='Heading4')

                    # Asegurarse de que la tabla tiene al menos una fila y una columna
                    if len(tabla) > 0 and len(tabla[0]) > 0:
                        doc_table = doc.add_table(rows=1, cols=len(tabla[0]))
                        
                        doc_table.style = 'Table Grid'

                        # encabezados de la tabla
                        for i, encabezado in enumerate(tabla[0]):
                            doc_table.cell(0, i).text = encabezado
                            paragraph = doc_table.cell(0, i).paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.runs
                            run[0].bold = True

                        # filas de datos a la tabla
                        for datos_fila in tabla[1:]:
                            row_cells = doc_table.add_row().cells
                            for i, valor_celda in enumerate(datos_fila):
                                row_cells[i].text = str(valor_celda)
                                paragraph = row_cells[i].paragraphs[0]
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
            except IndexError as e:
                print(f'Error al agregar tabla: {str(e)}')     

        # Manejo de las imágenes
        for imagen in data.get('graficas', []):
            if imagen.startswith('http://') or imagen.startswith('https://'):
                # URLs de imágenes: descargar y agregar a Word
                try:
                    respuesta = requests.get(imagen)
                    image_stream = BytesIO(respuesta.content)
                    doc.add_picture(image_stream, width=Inches(6))
                except Exception as e:
                    print(f'Error al descargar/agregar imagen desde URL: {str(e)}')
            else:
                try:
                    image_data = base64.b64decode(imagen.split(',')[1])
                    image_stream = BytesIO(image_data)
                    doc.add_picture(image_stream, width=Inches(6))
                except Exception as e:
                    print(f'Error al agregar imagen en base64: {str(e)}')

        # Guarda y envia el documento de Word
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        response = HttpResponse(doc_io.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Reporte_Análisis.docx"'
        return response

    except json.JSONDecodeError as e:
        return JsonResponse({'error': f'Error al decodificar JSON: {str(e)}'}, status=400)
    except Exception as e:
        return JsonResponse({'error': f'Error al procesar la solicitud: {str(e)}'}, status=500)


def reporte_regresion_linear(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    # Elimina los datos del reporte de la sesión después de usarlos
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_regresion_linear.html', {
        'file_name': file_name,
        'reporte': reporte_data
    })

def reporte_regresion_logistica(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_regresion_logistica.html', {
        'file_name': file_name,
        'reporte': reporte_data
    })


def reporte_tendencia_central(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_tendencia_central.html', {
        'file_name': file_name,
        'reporte_data': reporte_data
    })


def reporte_variabilidad(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_variabilidad.html', {
        'file_name': file_name,
        'reporte_data': reporte_data
    })

def reporte_distribucion_frecuencias(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_distribucion_frecuencias.html', {
        'file_name': file_name,
        'reporte_data': reporte_data
    })



def reporte_correlacion_pearson(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_correlacion_pearson.html', {
        'file_name': file_name,
        'reporte_data': reporte_data  
    })

def reporte_correlacion_spearman(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_correlacion_spearman.html', {
        'file_name': file_name,
        'reporte_data': reporte_data 
    })


def reporte_kmeans_clustering(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_kmeans_clustering.html', {
        'file_name': file_name,
        'reporte_data': reporte_data 
    })

def reporte_arbol_decision_clasificacion(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_arbol_decision_clasificacion.html', {
        'file_name': file_name,
        'reporte_data': reporte_data  
    })




def reporte_intervalo_confianza(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_intervalo_confianza.html', {
        'file_name': file_name,
        'reporte_data': reporte_data,
    })


def reporte_prueba_hipotesis(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_prueba_hipotesis.html', {
        'file_name': file_name,
        'reporte_data': reporte_data,
    })


def reporte_perceptron(request, file_name):
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)

    del request.session['reporte_data']
    

    return render(request, 'report/reporte_perceptron.html', {
        'file_name': file_name,
        'reporte_data': reporte_data,
    })
